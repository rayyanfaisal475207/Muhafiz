"""
Phase 7, Module 7.4 (item 4) — docx_loader._iter_blocks's O(n^2) rescan.

_iter_blocks used to re-scan word_doc.paragraphs/.tables from scratch for
every body child element to find the one matching by `is` identity — a
linear lookup dict built once, keyed by element id(), replaces that
without changing what's yielded (same objects, same order).
"""
from docx import Document as DocxDocument

from src.ingestion.loaders.docx_loader import _iter_blocks


def _build_docx(tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("First paragraph", style="Heading 1")
    doc.add_paragraph("Second paragraph, normal style")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "cell-a"
    table.cell(0, 1).text = "cell-b"
    doc.add_paragraph("Third paragraph, after the table")

    path = tmp_path / "sample.docx"
    doc.save(path)
    return DocxDocument(path)


def test_iter_blocks_preserves_document_order_across_paragraphs_and_tables(tmp_path):
    word_doc = _build_docx(tmp_path)

    blocks = list(_iter_blocks(word_doc))

    assert [b["type"] for b in blocks] == ["paragraph", "paragraph", "table", "paragraph"]
    assert blocks[0]["text"] == "First paragraph"
    assert blocks[0]["style"] == "Heading 1"
    assert blocks[1]["text"] == "Second paragraph, normal style"
    assert blocks[3]["text"] == "Third paragraph, after the table"


def test_iter_blocks_table_block_carries_the_real_table_object(tmp_path):
    word_doc = _build_docx(tmp_path)

    blocks = list(_iter_blocks(word_doc))
    table_block = next(b for b in blocks if b["type"] == "table")

    assert table_block["table"].cell(0, 0).text == "cell-a"
    assert table_block["table"].cell(0, 1).text == "cell-b"


def test_iter_blocks_matches_each_element_to_the_correct_object_not_just_any(tmp_path):
    """Regression guard for the O(n^2) fix specifically: with several
    paragraphs of the same style, each yielded block must correspond to
    the SAME element it was matched from, not an arbitrary same-style one
    (a bug a naive id()-based lookup with colliding keys could introduce)."""
    doc = DocxDocument()
    for i in range(6):
        doc.add_paragraph(f"Paragraph number {i}", style="Normal")
    path = tmp_path / "many.docx"
    doc.save(path)
    word_doc = DocxDocument(path)

    blocks = list(_iter_blocks(word_doc))

    assert [b["text"] for b in blocks] == [f"Paragraph number {i}" for i in range(6)]
